import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

mydoc = docx.Document()
mydoc.add_paragraph("My name is hashem")
mydoc.add_paragraph("i am the king")


Paragraph = mydoc.add_paragraph("السلام عليكم ورحمة الله وبركاته")
Paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

third_paragrapgh = mydoc.add_paragraph("this is my world.")
third_paragrapgh.add_run(" yes i am")

mydoc.add_heading("My World" , 0)
mydoc.add_heading("My World" , 1)
mydoc.add_heading("My World" , 2)
mydoc.add_heading("My World" , 3)
mydoc.add_heading("My World" , 4)
mydoc.add_heading("My World" , 5)

mydoc.paragraphs[0].style = mydoc.styles['Heading 1']


print(mydoc.paragraphs[0].style)
print(mydoc.paragraphs[5].style)
print(mydoc.paragraphs[4].style)

mydoc.save('write.docx')